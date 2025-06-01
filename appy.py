import streamlit as st
import os
import tempfile
import shutil
import docx
import pandas as pd
import re
from datetime import datetime
from pathlib import Path
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill
import xlwings as xw

# —————————————————————————————
# Page Layout & Header
# —————————————————————————————
st.set_page_config(
    page_title="PRL Timesheet Processor",
    layout="centered",
    initial_sidebar_state="expanded",
)

st.markdown("""
# 🗂️ PRL Timesheet Processor  
Upload one or more `.docx` timesheet files and click **“Process Timesheets”**.  
This app generates a consolidated Excel report (with formulas, a pivot table, and a weekly summary).  
**Requires Microsoft Excel installed** (so that `xlwings` can drive Excel).
""")

# —————————————————————————————
# Sidebar: Editable Rates & Overtime Rules
# —————————————————————————————
st.sidebar.header("⚙️ Edit Rates & Overtime")

# 1. Default hourly rate
default_rate = st.sidebar.number_input(
    "Default Rate (£/hr)", 
    min_value=0.0, 
    value=15.0, 
    step=0.5,
    help="If a name has no custom rate, this rate will be used."
)

# 2. Overtime threshold (total hours per week before a special multiplier applies)
overtime_threshold = st.sidebar.number_input(
    "Overtime threshold (hrs/week)", 
    min_value=0, 
    value=50, 
    step=1,
    help="After this many total hours in a week, an overtime multiplier applies."
)

# 3. Saturday and Sunday multipliers
saturday_multiplier = st.sidebar.number_input(
    "Saturday multiplier", 
    min_value=1.0, 
    value=1.5, 
    step=0.1,
    help="Multiply Saturday hours by this factor."
)
sunday_multiplier = st.sidebar.number_input(
    "Sunday multiplier", 
    min_value=1.0, 
    value=1.75, 
    step=0.1,
    help="Multiply Sunday hours by this factor."
)
overtime_multiplier = st.sidebar.number_input(
    "Overtime multiplier", 
    min_value=1.0, 
    value=1.25, 
    step=0.05,
    help="For hours beyond the weekly threshold, multiply by this factor (applied after day multipliers)."
)

st.sidebar.markdown("---")
st.sidebar.subheader("Custom Rates by Name")

# 4. Custom rates table (prepopulate with known names or add new)
custom_rates = {}
# You can pre-list known employees, or allow free-form entries:
prepop_names = ["Aaron Hall", "Finley Mc", "Andrew Burke"]
for name in prepop_names:
    rate_val = st.sidebar.number_input(
        f"{name} (£/hr)", 
        min_value=0.0, 
        value=0.0, 
        step=0.5, 
        key=name
    )
    if rate_val > 0:
        custom_rates[name] = rate_val

# Allow adding a new custom rate:
new_name = st.sidebar.text_input("Add custom name")
if new_name:
    new_rate = st.sidebar.number_input(
        f"{new_name} (£/hr)", 
        min_value=0.0, 
        value=0.0, 
        step=0.5, 
        key="new_" + new_name
    )
    if new_rate > 0:
        custom_rates[new_name] = new_rate

st.sidebar.caption("If the rate is zero, the default rate is used.")

# —————————————————————————————
# Functions: Extraction & Pay Calculation
# —————————————————————————————
def calculate_pay(name, daily_data):
    """
    Calculate total pay using:
    - Custom rate if provided, otherwise default_rate.
    - Day multipliers for Saturday/Sunday.
    - Overtime multiplier if weekly total exceeds threshold.
    """
    # Determine base hourly rate
    rate = custom_rates.get(name, default_rate)

    total_hours = 0.0
    pay = 0.0

    # 1. First pass: apply day multipliers and accumulate hours
    for day in daily_data:
        hours = day["hours"]
        weekday = day["weekday"]
        multiplier = 1.0
        if weekday == "Saturday":
            multiplier = saturday_multiplier
        elif weekday == "Sunday":
            multiplier = sunday_multiplier
        pay += hours * rate * multiplier
        total_hours += hours

    # 2. If total_hours exceed overtime_threshold, apply overtime multiplier to the excess hours
    if total_hours > overtime_threshold:
        excess = total_hours - overtime_threshold
        # Calculate pay for excess hours at (base rate * overtime_multiplier) minus what was already counted
        # For simplicity, assume excess hours are paid at base rate * overtime_multiplier (no day multiplier on excess).
        pay += excess * rate * (overtime_multiplier - 1)

    return round(total_hours, 2), round(pay, 2), rate

def extract_timesheet_data(doc_path):
    """
    Parse the .docx file to extract:
    - Name (default to filename if not found)
    - Client (from same line as "Client")
    - Site Address
    - Daily hours list for calculate_pay
    - Date range, etc.
    """
    doc = docx.Document(doc_path)
    name, client, site = "", "", ""
    daily_data = []

    # 1) Extract client and possibly name from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if "Client" in text and not client:
            match = re.search(r"Client\s*[:\t\-\s]*(.+)", text, re.IGNORECASE)
            if match:
                client = match.group(1).strip()

        if not name and re.fullmatch(r"[A-Z\s]{5,}", text) and "PRL" not in text:
            if len(text.split()) >= 2:
                name = text.title()

        if not name and "print name" in text.lower():
            match = re.search(r"print name[:\-\s]*(.+)", text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()

        if "Site Address" in text and not site:
            match = re.search(r"Site Address\s*[:\-\s]*(.+)", text, re.IGNORECASE)
            if match:
                site = match.group(1).strip()

    # 2) Look in tables for fallback
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for cell in cells:
                ct = cell.text.strip()
                if not site and "Site Address" in ct:
                    m2 = re.search(r"Site Address\s*[:\-\s]*(.+)", ct, re.IGNORECASE)
                    if m2:
                        site = m2.group(1).strip()

                if not name:
                    for line in ct.splitlines():
                        if re.fullmatch(r"[A-Z\s]{5,}", line) and "PRL" not in line:
                            if len(line.split()) >= 2:
                                name = line.title()

            # Extract daily hours
            if len(cells) >= 5:
                hrs = cells[4].text.strip()
                date_text = cells[0].text.strip()
                day_text = cells[1].text.strip()
                if hrs and hrs not in ["-", "–", "—"] and day_text:
                    try:
                        daily_data.append({"weekday": day_text, "hours": float(hrs)})
                    except:
                        pass

    total_hours, pay_amount, used_rate = calculate_pay(name or "", daily_data)

    # Compute date range
    date_list = []
    for table in doc.tables:
        for row in table.rows:
            dt = row.cells[0].text.strip()
            if re.match(r"\d{2}\.\d{2}\.\d{4}", dt):
                try:
                    date_list.append(datetime.strptime(dt, "%d.%m.%Y"))
                except:
                    pass
    date_range = ""
    if date_list:
        date_list = sorted(date_list)
        date_range = f"{date_list[0].strftime('%d.%m.%Y')}–{date_list[-1].strftime('%d.%m.%Y')}"

    # If name is blank, default to filename stem
    if not name:
        name = Path(doc_path).stem

    return {
        "Name": name,
        "Client": client,
        "Site Address": site,
        "Total Hours": total_hours,
        "Rate (£)": used_rate,
        "Calculated Pay (£)": pay_amount,
        "Printed Name": name,
        "Date Range": date_range,
        "File Name": Path(doc_path).name,
        "Extracted On": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

def process_folder(folder_path):
    results = []
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(".docx"):
            full_path = os.path.join(folder_path, filename)
            try:
                results.append(extract_timesheet_data(full_path))
            except Exception as e:
                st.error(f"Failed to process {filename}: {e}")
    return results

def autofit_and_style(wb_path):
    wb = load_workbook(wb_path)
    ws = wb["Timesheets"]
    for col in ws.columns:
        max_len = 0
        letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[letter].width = max_len + 2

    header_font = Font(bold=True)
    header_fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill

    wb.save(wb_path)

def apply_excel_formula_to_pay_column(file_path):
    """
    Note: Now that we calculate pay_amount in Python,
    the "Calculated Pay (£)" column contains the final value already.
    We still leave this here if you want Excel formulas instead.
    """
    wb = load_workbook(file_path)
    ws = wb["Timesheets"]
    headers = {c.value: i + 1 for i, c in enumerate(ws[1])}
    total_col = headers.get("Total Hours")
    rate_col = headers.get("Rate (£)")
    pay_col = headers.get("Calculated Pay (£)")
    if not all([total_col, rate_col, pay_col]):
        st.error("Missing required columns to apply pay formula.")
        return
    for r in range(2, ws.max_row + 1):
        hcell = f"{get_column_letter(total_col)}{r}"
        rcell = f"{get_column_letter(rate_col)}{r}"
        formula = f"={hcell}*{rcell}"
        ws.cell(row=r, column=pay_col).value = formula
    wb.save(file_path)

def create_pivot_table(wb_path):
    app = xw.App(visible=False)
    wb = app.books.open(wb_path)
    sht = wb.sheets["Timesheets"]
    if "Pivot Report" in [s.name for s in wb.sheets]:
        pivot_sht = wb.sheets["Pivot Report"]
        pivot_sht.clear()
    else:
        pivot_sht = wb.sheets.add("Pivot Report")

    last_row = sht.range("A1").end("down").row
    last_col = sht.range("A1").end("right").column
    source = sht.range((1, 1), (last_row, last_col))

    pivot = pivot_sht.api.PivotTableWizard(
        SourceType=1,
        SourceData=source.api,
        TableDestination=pivot_sht.range("A5").api,
        TableName="PivotTimesheet"
    )
    pivot.PivotFields("Name").Orientation = 1
    pivot.AddDataField(
        pivot.PivotFields("Calculated Pay (£)"),
        "Total Pay",
        -4157  # xlSum
    )
    pivot.PivotFields("Name").AutoSort(2, "Total Pay")
    pivot_sht.autofit()
    wb.save()
    wb.close()
    app.quit()

def get_calculated_pay_summary_with_xlwings(excel_path):
    app = xw.App(visible=False)
    wb = app.books.open(excel_path)
    sht = wb.sheets["Timesheets"]
    vals = sht.range("A1").expand().value
    headers = vals[0]
    rows = vals[1:]
    name_idx = headers.index("Name")
    pay_idx = headers.index("Calculated Pay (£)")
    summary = {}
    for row in rows:
        nm = row[name_idx]
        pay = row[pay_idx]
        if nm and isinstance(pay, (int, float)):
            summary[nm] = summary.get(nm, 0) + pay
    wb.close()
    app.quit()
    return summary

def write_weekly_summary_to_excel(summary_data, excel_path):
    wb = load_workbook(excel_path)
    if "Weekly Summary" not in wb.sheetnames:
        wb.create_sheet("Weekly Summary")
    ws = wb["Weekly Summary"]
    ws.delete_rows(1, ws.max_row)
    ws.append(["Name", "Total Pay (£)"])
    for nm, tot in summary_data.items():
        ws.append([nm, round(tot, 2)])
    wb.save(excel_path)

# —————————————————————————————
# Main UI: File Uploader & Process Button
# —————————————————————————————
uploaded_files = st.file_uploader(
    "Upload Word Timesheets (.docx)", 
    type="docx", 
    accept_multiple_files=True
)

if st.button("Process Timesheets") and uploaded_files:
    with st.spinner("Processing… this may take a moment while Excel is driven via xlwings."):
        temp_dir = tempfile.mkdtemp(prefix="timesheet_")
        try:
            # 1. Save uploads
            for f in uploaded_files:
                save_path = os.path.join(temp_dir, f.name)
                with open(save_path, "wb") as out_file:
                    out_file.write(f.getbuffer())

            # 2. Parse & calculate (Python‐side)
            results = process_folder(temp_dir)
            df = pd.DataFrame(results)

            # 3. Show a preview of parsed data
            st.markdown("### 📋 Parsed Timesheet Data")
            st.dataframe(df)

            # 4. Show an in‐browser summary
            summary_df = df.groupby("Name")[["Total Hours", "Calculated Pay (£)"]].sum().reset_index()
            st.markdown("### 💰 In‐Browser Pay Summary")
            st.dataframe(summary_df.rename(columns={"Calculated Pay (£)": "Total Pay (£)"}))

            # 5. Write to Excel
            output_file = os.path.join(temp_dir, "All_Timesheets_Combined.xlsx")
            df.to_excel(output_file, sheet_name="Timesheets", index=False)

            # 6. Style & pivot & summary within Excel
            autofit_and_style(output_file)
            apply_excel_formula_to_pay_column(output_file)
            create_pivot_table(output_file)
            excel_summary = get_calculated_pay_summary_with_xlwings(output_file)
            write_weekly_summary_to_excel(excel_summary, output_file)

            # 7. Provide download buttons
            with open(output_file, "rb") as file_bytes:
                data = file_bytes.read()

            st.success("✅ All done! Download your files below:")
            col1, col2 = st.columns(2)
            col1.download_button(
                label="⬇️ Download Excel Report",
                data=data,
                file_name="All_Timesheets_Combined.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            # Also offer CSV
            csv_buffer = df.to_csv(index=False).encode("utf-8")
            col2.download_button(
                label="⬇️ Download Raw CSV",
                data=csv_buffer,
                file_name="timesheet_data.csv",
                mime="text/csv"
            )
        except Exception as exc:
            st.error(f"Processing error: {exc}")
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

if not uploaded_files:
    st.info("Select one or more `.docx` files above and adjust rates in the sidebar before processing.")
