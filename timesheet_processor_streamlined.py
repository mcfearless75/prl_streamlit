import streamlit as st
import pandas as pd
import os
import docx
import re
from datetime import datetime
from io import BytesIO

# Custom hourly rates dictionary
custom_rates = {
    "Aaron Hall": 15.0,
    "Finley Mc": 18.0,
    "Andrew Burke": 16.5
}
default_rate = 15.0

# Overtime multipliers
OT_RULES = {
    "Saturday": 1.5,
    "Sunday": 1.75,
    "Over50": 1.5
}

def extract_timesheet_data(doc_file):
    doc = docx.Document(doc_file)
    name = client = site = printed_name = ""
    daily_data = []
    date_list = []

    for para in doc.paragraphs:
        if "Print Name" in para.text:
            match = re.search(r"Print Name[:\-]?\s*(.+)", para.text, re.IGNORECASE)
            if match:
                printed_name = match.group(1).strip()

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for cell in cells:
                txt = cell.text.strip()
                if "Client" in txt and not client:
                    match = re.search(r"Client[:\-]?\s*(.+)", txt, re.IGNORECASE)
                    if match:
                        client = match.group(1).strip()
                if "Site Address" in txt and not site:
                    match = re.search(r"Site Address[:\-]?\s*(.+)", txt, re.IGNORECASE)
                    if match:
                        site = match.group(1).strip()
                if not name and re.fullmatch(r"[A-Z\s]{5,}", txt) and "PRL" not in txt:
                    name = txt.title()

            if len(cells) >= 5:
                date_text = cells[0].text.strip()
                day = cells[1].text.strip()
                hrs = cells[4].text.strip()
                if re.match(r"\d{2}\.\d{2}\.\d{4}", date_text):
                    try:
                        date_obj = datetime.strptime(date_text, "%d.%m.%Y")
                        date_list.append(date_obj)
                        if hrs and hrs not in ["-", "–", ""]:
                            daily_data.append({
                                "date": date_obj,
                                "weekday": day,
                                "hours": float(hrs)
                            })
                    except:
                        continue

    if not name and printed_name:
        name = printed_name

    total_hours = sum(d['hours'] for d in daily_data)
    weekend_hours = sum(d['hours'] for d in daily_data if d['weekday'] in ["Saturday", "Sunday"])
    over_50 = max(total_hours - 50, 0)
    rate = custom_rates.get(name, default_rate)

    regular_hours = total_hours - weekend_hours - over_50
    regular_pay = regular_hours * rate
    saturday_pay = sum(d['hours'] for d in daily_data if d['weekday'] == "Saturday") * rate * OT_RULES["Saturday"]
    sunday_pay = sum(d['hours'] for d in daily_data if d['weekday'] == "Sunday") * rate * OT_RULES["Sunday"]
    over50_pay = over_50 * rate * OT_RULES["Over50"]
    total_pay = regular_pay + saturday_pay + sunday_pay + over50_pay

    date_range = f"{min(date_list).strftime('%d.%m.%Y')} – {max(date_list).strftime('%d.%m.%Y')}" if date_list else ""

    return {
        "Name": name,
        "Client": client,
        "Site": site,
        "Date Range": date_range,
        "Total Hours": round(total_hours, 2),
        "Base Rate (£)": rate,
        "Over 50 Hours (hrs)": round(over_50, 2),
        "Saturday OT (hrs)": round(sum(d['hours'] for d in daily_data if d['weekday'] == "Saturday"), 2),
        "Sunday OT (hrs)": round(sum(d['hours'] for d in daily_data if d['weekday'] == "Sunday"), 2),
        "Calculated Pay (£)": round(total_pay, 2)
    }

st.title("📄 PRL Timesheet Processor")

uploaded_files = st.file_uploader("Upload one or more .docx timesheets", type=["docx"], accept_multiple_files=True)

if uploaded_files:
    rows = []
    for f in uploaded_files:
        rows.append(extract_timesheet_data(f))

    df = pd.DataFrame(rows)
    st.dataframe(df)

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name="Timesheets", index=False)

    st.download_button(
        label="📥 Download Excel Summary",
        data=buffer.getvalue(),
        file_name="PRL_Timesheet_Summary.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )